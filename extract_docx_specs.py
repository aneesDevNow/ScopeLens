#!/usr/bin/env python3
"""Extract fonts, sizes, images, and styling from AI_Final Report.docx"""

import zipfile
import os
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt, Emu
from collections import defaultdict

DOCX_PATH = "/Users/manees/Mac_sync/ScopeLens/AI_Final Report.docx"
EXTRACT_DIR = "/Users/manees/Mac_sync/ScopeLens/docx_extracted"

# ─── 1. Extract embedded images ───
print("=" * 60)
print("1. EMBEDDED IMAGES / ICONS")
print("=" * 60)

os.makedirs(EXTRACT_DIR, exist_ok=True)
with zipfile.ZipFile(DOCX_PATH, 'r') as z:
    media_files = [f for f in z.namelist() if f.startswith('word/media/')]
    if media_files:
        for mf in media_files:
            fname = os.path.basename(mf)
            out_path = os.path.join(EXTRACT_DIR, fname)
            with z.open(mf) as src, open(out_path, 'wb') as dst:
                dst.write(src.read())
            size = os.path.getsize(out_path)
            print(f"  📎 {fname} ({size:,} bytes) → {out_path}")
    else:
        print("  No embedded media found")

# ─── 2. Extract fonts and sizes ───
print("\n" + "=" * 60)
print("2. FONTS & SIZES USED")
print("=" * 60)

doc = Document(DOCX_PATH)
font_usage = defaultdict(set)  # font_name -> set of sizes
style_info = {}

# Check document default font
if doc.styles.default(1):  # WD_STYLE_TYPE.PARAGRAPH = 1
    default_style = doc.styles.default(1)
    if default_style.font:
        print(f"\n  Default paragraph font: {default_style.font.name}, size: {default_style.font.size}")

# Scan all styles
print("\n  --- Document Styles ---")
for style in doc.styles:
    if not hasattr(style, 'font') or not style.font:
        continue
    if style.font.name or style.font.size:
        fname = style.font.name or "(inherited)"
        fsize = f"{style.font.size.pt}pt" if style.font.size else "(inherited)"
        fbold = style.font.bold
        fitalic = style.font.italic
        fcolor = None
        if style.font.color and style.font.color.rgb:
            fcolor = f"#{style.font.color.rgb}"
        print(f"  Style '{style.name}': font={fname}, size={fsize}, bold={fbold}, italic={fitalic}, color={fcolor}")

# Scan every paragraph and run
print("\n  --- Paragraph-level Analysis ---")
fonts_found = defaultdict(lambda: {"sizes": set(), "bold_sizes": set(), "colors": set(), "samples": []})

for i, para in enumerate(doc.paragraphs):
    if not para.text.strip():
        continue
    
    para_style = para.style.name if para.style else "None"
    
    for run in para.runs:
        font_name = run.font.name or (para.style.font.name if para.style and para.style.font else "Default")
        font_size = run.font.size
        is_bold = run.font.bold or (para.style.font.bold if para.style and para.style.font else False)
        is_italic = run.font.italic or (para.style.font.italic if para.style and para.style.font else False)
        
        color_str = None
        if run.font.color and run.font.color.rgb:
            color_str = f"#{run.font.color.rgb}"
        
        size_str = f"{font_size.pt}pt" if font_size else "inherited"
        
        entry = fonts_found[font_name]
        if font_size:
            if is_bold:
                entry["bold_sizes"].add(font_size.pt)
            else:
                entry["sizes"].add(font_size.pt)
        if color_str:
            entry["colors"].add(color_str)
        
        # Keep first 2 samples
        if len(entry["samples"]) < 2:
            sample = run.text[:60].strip()
            if sample:
                entry["samples"].append(f"[{size_str}, bold={is_bold}, italic={is_italic}, color={color_str}] \"{sample}\"")

for font_name, info in sorted(fonts_found.items()):
    print(f"\n  Font: {font_name}")
    if info["sizes"]:
        print(f"    Regular sizes: {sorted(info['sizes'])}")
    if info["bold_sizes"]:
        print(f"    Bold sizes: {sorted(info['bold_sizes'])}")
    if info["colors"]:
        print(f"    Colors: {sorted(info['colors'])}")
    for s in info["samples"]:
        print(f"    Sample: {s}")

# ─── 3. Paragraph styles breakdown ───
print("\n" + "=" * 60)
print("3. PARAGRAPH STYLES BREAKDOWN")
print("=" * 60)

style_counts = defaultdict(int)
for para in doc.paragraphs:
    if para.text.strip():
        style_name = para.style.name if para.style else "None"
        style_counts[style_name] += 1

for style_name, count in sorted(style_counts.items(), key=lambda x: -x[1]):
    print(f"  {style_name}: {count} paragraphs")

# ─── 4. Table analysis ───
print("\n" + "=" * 60)
print("4. TABLES")
print("=" * 60)
print(f"  Total tables: {len(doc.tables)}")
for i, table in enumerate(doc.tables):
    print(f"  Table {i+1}: {len(table.rows)} rows x {len(table.columns)} cols")
    for ri, row in enumerate(table.rows):
        if ri < 3:  # first 3 rows
            cells = [cell.text[:30] for cell in row.cells]
            print(f"    Row {ri}: {cells}")

# ─── 5. Section/page setup ───
print("\n" + "=" * 60)
print("5. PAGE SETUP")
print("=" * 60)
for i, section in enumerate(doc.sections):
    print(f"  Section {i+1}:")
    print(f"    Page width: {section.page_width.inches:.2f} in")
    print(f"    Page height: {section.page_height.inches:.2f} in")
    print(f"    Top margin: {section.top_margin.inches:.2f} in")
    print(f"    Bottom margin: {section.bottom_margin.inches:.2f} in")
    print(f"    Left margin: {section.left_margin.inches:.2f} in")
    print(f"    Right margin: {section.right_margin.inches:.2f} in")

# ─── 6. Extract raw XML for deeper analysis ───
print("\n" + "=" * 60)
print("6. EMBEDDED RELATIONSHIPS (images, links)")  
print("=" * 60)

with zipfile.ZipFile(DOCX_PATH, 'r') as z:
    # Check for relationships
    if 'word/_rels/document.xml.rels' in z.namelist():
        rels_xml = z.read('word/_rels/document.xml.rels').decode('utf-8')
        root = ET.fromstring(rels_xml)
        ns = {'r': 'http://schemas.openxmlformats.org/package/2006/relationships'}
        for rel in root.findall('r:Relationship', ns):
            rel_type = rel.get('Type', '').split('/')[-1]
            target = rel.get('Target', '')
            if rel_type in ('image', 'hyperlink'):
                print(f"  {rel_type}: {target}")

print("\n✅ Extraction complete!")
print(f"📁 Images saved to: {EXTRACT_DIR}")
